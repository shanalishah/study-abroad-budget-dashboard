import io
import re
import json
import hashlib
from pathlib import Path
from datetime import datetime, timezone
from zoneinfo import ZoneInfo

import pandas as pd
import streamlit as st

try:
    import yaml  # optional
except Exception:
    yaml = None

from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch

# =========================================================
# STREAMLIT CONFIG
# =========================================================
st.set_page_config(page_title="Education Abroad Budget Dashboard", layout="wide")
st.title("Education Abroad Budget Dashboard")

# =========================================================
# CONSTANTS
# =========================================================
SUPPORTED_COST_EXTS = [".tsv", ".txt", ".csv", ".yaml", ".yml", ".json"]

# statuses we consider "approved" for the dashboard
APPROVED_STATUSES = [
    "Committed",
    "Accepted",
    "Permission to Study Abroad: Granted",
    "Provisional Permission",
]

# and their priority for de-dup
STATUS_PRIORITY = {
    "Committed": 4,
    "Accepted": 3,
    "Permission to Study Abroad: Granted": 2,
    "Provisional Permission": 1,
}
def status_rank(s: str) -> int:
    return STATUS_PRIORITY.get(str(s), 0)

# =========================================================
# PATH / FILE HELPERS
# =========================================================
def _candidate_dirs():
    raw_dirs = [
        Path("./data"),
        Path("."),
        Path("/mount/src/study-abroad-budget-dashboard"),
        Path("/mount/src/study-abroad-budget-dashboard/data"),
    ]
    out, seen = [], set()
    for d in raw_dirs:
        try:
            dp = d.resolve()
        except Exception:
            dp = d
        if dp.exists() and str(dp) not in seen:
            seen.add(str(dp))
            out.append(dp)
    return out

def file_md5(path: Path) -> str:
    if not path or not path.exists():
        return ""
    return hashlib.md5(path.read_bytes()).hexdigest()

# filenames like: spring_2026_budget-2025-10-30-10_30_00.csv
def parse_new_budget_name(stem: str) -> datetime | None:
    m = re.search(r"(\d{4})-(\d{2})-(\d{2})-(\d{2})_(\d{2})_(\d{2})$", stem)
    if not m:
        return None
    y, mo, d, hh, mm, ss = map(int, m.groups())
    try:
        return datetime(y, mo, d, hh, mm, ss)
    except ValueError:
        return None

def find_latest_app_file() -> Path | None:
    data_dir = Path("./data")
    if not data_dir.exists():
        return None

    candidates: list[tuple[Path, datetime, datetime]] = []
    for p in data_dir.glob("*.csv"):
        parsed = parse_new_budget_name(p.stem)
        if parsed is None:
            continue
        try:
            mtime = datetime.fromtimestamp(
                p.stat().st_mtime, tz=timezone.utc
            ).astimezone(ZoneInfo("America/New_York")).replace(tzinfo=None)
        except Exception:
            mtime = datetime.min
        candidates.append((p, parsed, mtime))

    if not candidates:
        return None

    candidates.sort(key=lambda t: (t[1], t[2]))
    return candidates[-1][0]

def parse_timestamp_label(stem: str) -> str:
    dt = parse_new_budget_name(stem)
    if dt is None:
        return stem
    et = dt.replace(tzinfo=ZoneInfo("America/New_York"))
    return et.strftime("%Y-%m-%d %H:%M:%S %Z")

# =========================================================
# GENERIC READERS
# =========================================================
def _read_csv_forgiving(path_or_buf, sep=",", dtype=str):
    encodings = ["utf-8", "utf-8-sig", "cp1252"]
    last_err = None
    for enc in encodings:
        try:
            return pd.read_csv(
                path_or_buf,
                sep=sep,
                dtype=dtype,
                engine="python",
                encoding=enc,
                encoding_errors="replace",
                on_bad_lines="skip",
            )
        except Exception as e:
            last_err = e
    raise last_err

def _clean_col(c: str) -> str:
    return (
        str(c)
        .replace("\ufeff", "")   # BOM
        .replace("\u00a0", " ")  # NBSP
        .strip()
    )

def read_any_apps_file(file_obj_or_path):
    """Read CSV or TXT/TSV and normalize headers/values."""
    if file_obj_or_path is None:
        return None

    name = getattr(file_obj_or_path, "name", str(file_obj_or_path))
    if name and name.lower().endswith((".txt", ".tsv")):
        df = _read_csv_forgiving(file_obj_or_path, sep="\t", dtype=str)
    else:
        df = _read_csv_forgiving(file_obj_or_path, sep=",", dtype=str)

    df.columns = [_clean_col(c) for c in df.columns]
    for c in df.columns:
        if df[c].dtype == object:
            df[c] = df[c].astype(str).str.replace("\u00a0", " ").str.strip()
    return df

# =========================================================
# COSTS
# =========================================================
def _read_text(path: Path) -> str:
    for enc in ["utf-8", "utf-8-sig", "cp1252"]:
        try:
            return path.read_text(encoding=enc, errors="replace")
        except Exception:
            continue
    return path.read_bytes().decode("utf-8", errors="replace")

@st.cache_data(show_spinner=False, ttl=0)
def load_costs_from_repo(path: Path, content_hash: str) -> pd.DataFrame:
    if not path or not path.exists():
        st.warning("No cost file found. Expected ProgramCost.tsv/.txt/.csv/.yaml/.yml/.json.")
        return pd.DataFrame(columns=["__Program", "__Cost"])

    ext = path.suffix.lower()

    if ext in [".tsv", ".txt"]:
        df = _read_csv_forgiving(path, sep="\t", dtype=str)
    elif ext == ".csv":
        df = _read_csv_forgiving(path, sep=",", dtype=str)
    elif ext in [".yaml", ".yml"]:
        if yaml is None:
            st.error("YAML support requires pyyaml.")
            return pd.DataFrame(columns=["__Program", "__Cost"])
        txt = _read_text(path)
        data = yaml.safe_load(txt) if txt else []
        df = pd.DataFrame(data)
    elif ext == ".json":
        txt = _read_text(path)
        data = json.loads(txt) if txt else []
        df = pd.DataFrame(data)
    else:
        st.error(f"Unsupported cost format: {ext}")
        return pd.DataFrame(columns=["__Program", "__Cost"])

    df.columns = [_clean_col(c) for c in df.columns]

    name_col = next((c for c in ["Program Name", "Program_Name", "program", "Program", "Name"] if c in df.columns), None)
    cost_col = next((c for c in ["$ Cost/Student", "Cost", "cost", "Cost per Student", "Cost_Student"] if c in df.columns), None)

    if name_col is None or cost_col is None:
        st.warning("Cost file must have Program Name + Cost column.")
        return pd.DataFrame(columns=["__Program", "__Cost"])

    df = df[~df[name_col].isna() & (df[name_col].astype(str).str.strip() != "")].copy()
    df["__Program"] = df[name_col].map(lambda x: coalesce_program_keys(x))
    df["__Cost"] = df[cost_col].map(normalize_currency)
    df = df[~df["__Program"].isna() & (df["__Program"].astype(str).str.strip() != "")]
    return df[["__Program", "__Cost"]]

def find_cost_file() -> Path | None:
    candidates = []
    for d in _candidate_dirs():
        for ext in SUPPORTED_COST_EXTS:
            p = d / f"ProgramCost{ext}"
            if p.exists():
                candidates.append(p)
    if not candidates:
        return None
    preference = {".tsv": 0, ".txt": 0, ".csv": 1, ".yaml": 2, ".yml": 2, ".json": 3}
    candidates.sort(key=lambda p: (preference.get(p.suffix.lower(), 9), str(p)))
    return candidates[0]

# =========================================================
# NORMALIZATION HELPERS
# =========================================================
def normalize_currency(s):
    if pd.isna(s):
        return pd.NA
    s = str(s).replace(",", "").replace("$", "").strip()
    if s == "" or s.lower() == "na":
        return pd.NA
    try:
        return float(s)
    except Exception:
        return pd.NA

def coalesce_program_keys(name):
    if pd.isna(name):
        return None
    s = str(name).strip()
    s = re.sub("[\u2013\u2014]", "-", s)
    s = s.replace("’", "'").replace("“", '"').replace("”", '"')
    s = re.sub(r"\s+", " ", s)
    return s

def resolve_columns(df: pd.DataFrame):
    expected = {
        "Program_Name": ["Program_Name", "Program Name", "Program"],
        "Application_Status": ["Application_Status", "Application Status", "Status"],
        "Program_Term": ["Program_Term", "Program Term", "Term"],
        "Program_Year": ["Program_Year", "Program Year", "Year"],
    }
    resolved = {}
    lower_map = {c.lower(): c for c in df.columns}
    for key, candidates in expected.items():
        found = None
        for c in candidates:
            if c in df.columns:
                found = c
                break
            lc = c.lower()
            if lc in lower_map:
                found = lower_map[lc]
                break
        if found:
            resolved[key] = found
    missing = [k for k in expected if k not in resolved]
    return resolved, missing

# =========================================================
# DEDUP BY STUDENT (priority + cost)
# =========================================================
def build_student_key(df: pd.DataFrame) -> pd.Series:
    candidates = [
        "UR ID", "UR_ID", "Student_ID", "Student Id", "ID",
        "User ID", "UserID", "User Id",
        "Email", "Email_Address", "Email Address", "Student Email", "User Email",
        "NetID", "Net Id",
    ]
    found = next((c for c in candidates if c in df.columns), None)
    if found:
        return df[found].astype(str).str.strip().str.lower().replace({"nan": ""})

    fn = df.columns[df.columns.str.lower().str.contains("first")].tolist()
    ln = df.columns[df.columns.str.lower().str.contains("last")].tolist()
    if fn and ln:
        return (
            df[fn[0]].astype(str).str.strip().str.lower()
            + "||"
            + df[ln[0]].astype(str).str.strip().str.lower()
        )

    return (
        df.get("Email", pd.Series([""] * len(df))).astype(str).str.lower()
        + "||"
        + df.get("__Program", pd.Series([""] * len(df))).astype(str).str.lower()
    )

def dedup_students_by_priority_then_cost(df_apps: pd.DataFrame, df_costs: pd.DataFrame) -> pd.DataFrame:
    if df_apps.empty:
        return df_apps

    costs = df_costs.rename(
        columns={"__Program": "__Program_cost_key", "__Cost": "__CostValue"}
    ).copy()

    tmp = df_apps.copy()
    tmp["__StudentKey"] = build_student_key(tmp)
    tmp["__Rank"] = tmp["__Status"].map(status_rank)

    tmp = tmp.merge(
        costs[["__Program_cost_key", "__CostValue"]],
        left_on="__Program",
        right_on="__Program_cost_key",
        how="left",
    )
    tmp["__CostValue"] = tmp["__CostValue"].fillna(0.0)

    # Highest rank first, then higher cost
    tmp = tmp.sort_values(["__Rank", "__CostValue"], ascending=[False, False], kind="stable")

    # keep best per student
    tmp = tmp.drop_duplicates(subset="__StudentKey", keep="first")

    return tmp.drop(columns=["__StudentKey", "__Rank", "__Program_cost_key", "__CostValue"], errors="ignore")

# =========================================================
# AGGREGATION
# =========================================================
def aggregate_by_program_status(df, costs):
    if df.empty:
        return pd.DataFrame(
            columns=[
                "Program",
                "Status",
                "Student Count",
                "Cost per Student (USD)",
                "Budget (USD)",
            ]
        )

    counts = (
        df.groupby(["__Program", "__Status"], dropna=False)
        .size()
        .reset_index(name="Student Count")
    )
    merged = counts.merge(costs, on="__Program", how="left")

    merged["Program"] = merged["__Program"]
    merged["Status"] = merged["__Status"]
    merged["Cost per Student (USD)"] = merged["__Cost"]
    merged["Budget (USD)"] = merged["Student Count"] * merged["__Cost"]

    return merged[
        ["Program", "Status", "Student Count", "Cost per Student (USD)", "Budget (USD)"]
    ]

def kpi_row(df):
    total_students = int(df["Student Count"].sum()) if not df.empty else 0
    total_budget = df["Budget (USD)"].sum(min_count=1) if not df.empty else pd.NA
    return total_students, total_budget

def fmt_money(x):
    return "-" if pd.isna(x) else f"${x:,.0f}"

# =========================================================
# SIDEBAR: FILE PICKER
# =========================================================
st.sidebar.header("Applications Data")
uploaded_apps = st.sidebar.file_uploader("Upload file (CSV or TXT)", type=["csv", "txt", "tsv"])

if "use_repo_default" not in st.session_state:
    st.session_state.use_repo_default = True

if uploaded_apps is not None:
    st.session_state.use_repo_default = False

if st.sidebar.button("Use latest repository file"):
    st.session_state.use_repo_default = True

DEFAULT_APPS_PATH = find_latest_app_file()
if st.session_state.use_repo_default:
    if DEFAULT_APPS_PATH:
        st.sidebar.success(f"Source: {DEFAULT_APPS_PATH.name} ({parse_timestamp_label(DEFAULT_APPS_PATH.stem)})")
    else:
        st.sidebar.warning("No application file detected in ./data.")
else:
    st.sidebar.info("Source: uploaded file (session only)")

# =========================================================
# LOAD COSTS
# =========================================================
COST_FILE_PATH = find_cost_file()
COSTS = load_costs_from_repo(COST_FILE_PATH, file_md5(COST_FILE_PATH))

# =========================================================
# LOAD APPLICATIONS
# =========================================================
apps_df = (
    read_any_apps_file(DEFAULT_APPS_PATH)
    if st.session_state.use_repo_default
    else read_any_apps_file(uploaded_apps)
)
if apps_df is None or apps_df.empty:
    st.stop()

# =========================================================
# RESOLVE REQUIRED COLUMNS
# =========================================================
resolved, missing = resolve_columns(apps_df)
if missing:
    # extra fallback based on lower-cased headers
    colmap = {c.lower().replace("\u00a0", " ").strip(): c for c in apps_df.columns}
    if "program name" in colmap and "Program_Name" not in resolved:
        resolved["Program_Name"] = colmap["program name"]
    if "application status" in colmap and "Application_Status" not in resolved:
        resolved["Application_Status"] = colmap["application status"]
    if "program term" in colmap and "Program_Term" not in resolved:
        resolved["Program_Term"] = colmap["program term"]
    if "program year" in colmap and "Program_Year" not in resolved:
        resolved["Program_Year"] = colmap["program year"]

    missing = [k for k in ["Program_Name", "Application_Status", "Program_Term", "Program_Year"] if k not in resolved]

if missing:
    st.error(
        "Missing required columns in applications data: "
        + ", ".join([m.replace("_", " ") for m in missing])
        + ". I accept flexible names like 'Program'/'Status' too."
    )
    st.write("DEBUG headers:", list(apps_df.columns))
    st.stop()

# =========================================================
# NORMALIZE KEY COLS
# =========================================================
apps_df["__Program"] = apps_df[resolved["Program_Name"]].map(coalesce_program_keys)
apps_df["__Status"] = apps_df[resolved["Application_Status"]].fillna("")
apps_df["__Term"] = apps_df[resolved.get("Program_Term", "")] if "Program_Term" in resolved else ""
apps_df["__Year"] = apps_df[resolved.get("Program_Year", "")] if "Program_Year" in resolved else ""

# =========================================================
# FILTERS
# =========================================================
st.sidebar.header("Filters")
term_opts = ["(all)"] + sorted([t for t in apps_df["__Term"].dropna().unique() if str(t).strip()])
year_opts = ["(all)"] + sorted([y for y in apps_df["__Year"].dropna().unique() if str(y).strip()])
term = st.sidebar.selectbox("Program term", options=term_opts, index=0)
year = st.sidebar.selectbox("Program year", options=year_opts, index=0)

filtered = apps_df.copy()
if term != "(all)":
    filtered = filtered[filtered["__Term"] == term]
if year != "(all)":
    filtered = filtered[filtered["__Year"] == year]

# =========================================================
# DEDUP (this is where priority kicks in)
# =========================================================
filtered = dedup_students_by_priority_then_cost(filtered, COSTS)

# =========================================================
# SPLIT COHORTS
# =========================================================
confirmed = filtered[filtered["__Status"].isin(APPROVED_STATUSES)].copy()
pending = filtered[~filtered["__Status"].isin(APPROVED_STATUSES)].copy()
confirmed["__Group"] = "Approved"
pending["__Group"] = "Pending"

# =========================================================
# AGGREGATION + KPIs
# =========================================================
agg_confirmed = aggregate_by_program_status(confirmed, COSTS)
agg_pending = aggregate_by_program_status(pending, COSTS)

conf_students, conf_budget = kpi_row(agg_confirmed)
pend_students, pend_budget = kpi_row(agg_pending)

total_students = conf_students + pend_students
total_budget = (0 if pd.isna(conf_budget) else conf_budget) + (0 if pd.isna(pend_budget) else pend_budget)

k1, k2, k3, k4, k5, k6 = st.columns(6)
with k1:
    st.metric("Approved Students", f"{conf_students:,}")
with k2:
    st.metric("Approved Budget Exposure", fmt_money(conf_budget))
with k3:
    st.metric("Pending Students", f"{pend_students:,}")
with k4:
    st.metric("Pending Budget Exposure", fmt_money(pend_budget))
with k5:
    st.metric("Total Students", f"{total_students:,}")
with k6:
    st.metric("Total Budget Exposure", fmt_money(total_budget))

# =========================================================
# TABLES
# =========================================================
st.subheader("Approved - Student Budget")
if agg_confirmed.empty:
    st.write("No records for the current filters.")
else:
    st.dataframe(agg_confirmed.sort_values(["Program", "Status"]), use_container_width=True)

st.subheader("Pending - Student Budget")
if agg_pending.empty:
    st.write("No records for the current filters.")
else:
    st.dataframe(agg_pending.sort_values(["Program", "Status"]), use_container_width=True)

# =========================================================
# STUDENT EXPORTS
# =========================================================
student_cols_preferred = []
for label in [
    "UR ID", "UR_ID", "Student_ID", "Student Id", "ID", "User ID",
    "First Name", "First_Name", "First", "User First Name",
    "Last Name", "Last_Name", "Last", "User Last Name",
    "Email", "Email_Address", "Email Address", "Student Email", "User Email",
    resolved["Program_Name"],
    resolved["Application_Status"],
    resolved.get("Program_Term", ""),
    resolved.get("Program_Year", ""),
]:
    if label and label in apps_df.columns and label not in student_cols_preferred:
        student_cols_preferred.append(label)

def build_student_export(df):
    if df.empty:
        return pd.DataFrame()
    cols = [c for c in student_cols_preferred if c in df.columns]
    export = df[cols].copy()
    export.columns = [c.replace("_", " ").strip() for c in export.columns]
    return export

confirmed_students_export = build_student_export(confirmed)
pending_students_export = build_student_export(pending)

combined_students_export = pd.DataFrame()
if not confirmed_students_export.empty:
    tmp = confirmed_students_export.copy()
    tmp.insert(0, "Cohort", "Approved")
    combined_students_export = pd.concat([combined_students_export, tmp], ignore_index=True)
if not pending_students_export.empty:
    tmp = pending_students_export.copy()
    tmp.insert(0, "Cohort", "Pending")
    combined_students_export = pd.concat([combined_students_export, tmp], ignore_index=True)

st.subheader("All Students - Combined List")
if combined_students_export.empty:
    st.write("No student records for the current filters.")
else:
    st.dataframe(combined_students_export, use_container_width=True)

st.subheader("Student Lists - Downloads")
c1, c2, c3 = st.columns(3)
with c1:
    if confirmed_students_export.empty:
        st.button("Download Approved (CSV)", disabled=True)
    else:
        st.download_button(
            "Download Approved (CSV)",
            data=confirmed_students_export.to_csv(index=False).encode("utf-8"),
            file_name="confirmed_approved_students.csv",
            mime="text/csv",
        )
with c2:
    if pending_students_export.empty:
        st.button("Download Pending (CSV)", disabled=True)
    else:
        st.download_button(
            "Download Pending (CSV)",
            data=pending_students_export.to_csv(index=False).encode("utf-8"),
            file_name="preliminary_pending_students.csv",
            mime="text/csv",
        )
with c3:
    if combined_students_export.empty:
        st.button("Download Combined Student List (CSV)", disabled=True)
    else:
        st.download_button(
            "Download Combined Student List (CSV)",
            data=combined_students_export.to_csv(index=False).encode("utf-8"),
            file_name="all_students_combined.csv",
            mime="text/csv",
        )

# =========================================================
# SUMMARY
# =========================================================
def totals_by(df):
    if df.empty:
        return pd.DataFrame(columns=["Status", "Students", "Budget"])
    t = (
        df.groupby("Status", dropna=False)[["Student Count", "Budget (USD)"]]
        .sum(min_count=1)
        .reset_index()
        .rename(columns={"Student Count": "Students", "Budget (USD)": "Budget"})
    )
    return t

summary = pd.concat(
    [totals_by(agg_confirmed), totals_by(agg_pending)],
    ignore_index=True,
)

st.subheader("Budget Summary by Status")
st.dataframe(summary, use_container_width=True)

# =========================================================
# DIAGNOSTICS
# =========================================================
all_agg = pd.concat([agg_confirmed, agg_pending], ignore_index=True)
missing_cost = sorted(
    all_agg[all_agg["Cost per Student (USD)"].isna()]["Program"]
    .dropna()
    .unique()
    .tolist()
)
if missing_cost:
    with st.expander("Programs Missing Cost Mapping"):
        st.write(missing_cost)

# =========================================================
# REPORTS
# =========================================================
def build_docx_report() -> bytes:
    doc = Document()
    doc.add_heading("Education Abroad Budget Snapshot", level=1)

    dt_label = datetime.now(ZoneInfo("America/New_York")).strftime("%Y-%m-%d %H:%M:%S %Z")
    src_label = (
        DEFAULT_APPS_PATH.name
        if (st.session_state.use_repo_default and DEFAULT_APPS_PATH)
        else "Uploaded file"
    )
    doc.add_paragraph(f"Generated: {dt_label}")
    doc.add_paragraph(f"Data source: {src_label}")

    doc.add_heading("Key Metrics", level=2)
    doc.add_paragraph(f"Confirmed Students: {conf_students:,}")
    doc.add_paragraph(f"Confirmed Budget Exposure: {fmt_money(conf_budget)}")
    doc.add_paragraph(f"Pending Students: {pend_students:,}")
    doc.add_paragraph(f"Pending Budget Exposure: {fmt_money(pend_budget)}")
    doc.add_paragraph(f"Total Students: {total_students:,}")
    doc.add_paragraph(f"Total Budget Exposure: {fmt_money(total_budget)}")

    doc.add_heading("Budget Summary by Status", level=2)
    if not summary.empty:
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Status", "Students", "Budget"
        for _, r in summary.iterrows():
            row = table.add_row().cells
            row[0].text = str(r["Status"])
            row[1].text = f'{int(r["Students"]):,}'
            val = r["Budget"]
            row[2].text = "-" if pd.isna(val) else f"${val:,.0f}"
    else:
        doc.add_paragraph("No summary available for current filters.")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def build_pdf_report() -> bytes:
    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=LETTER)
    width, height = LETTER
    y = height - 1.0 * inch

    c.setFont("Helvetica-Bold", 14)
    c.drawString(1.0 * inch, y, "Education Abroad Budget Snapshot")
    y -= 0.3 * inch
    c.setFont("Helvetica", 10)

    dt_label = datetime.now(ZoneInfo("America/New_York")).strftime("%Y-%m-%d %H:%M:%S %Z")
    src_label = (
        DEFAULT_APPS_PATH.name
        if (st.session_state.use_repo_default and DEFAULT_APPS_PATH)
        else "Uploaded file"
    )
    c.drawString(1.0 * inch, y, f"Generated: {dt_label}")
    y -= 0.2 * inch
    c.drawString(1.0 * inch, y, f"Data source: {src_label}")

    y -= 0.35 * inch
    c.setFont("Helvetica-Bold", 12)
    c.drawString(1.0 * inch, y, "Key Metrics")
    c.setFont("Helvetica", 10)
    y -= 0.25 * inch
    lines = [
        f"Confirmed Students: {conf_students:,}",
        f"Confirmed Budget Exposure: {fmt_money(conf_budget)}",
        f"Pending Students: {pend_students:,}",
        f"Pending Budget Exposure: {fmt_money(pend_budget)}",
        f"Total Students: {total_students:,}",
        f"Total Budget Exposure: {fmt_money(total_budget)}",
    ]
    for line in lines:
        c.drawString(1.0 * inch, y, line)
        y -= 0.2 * inch
        if y < 1.0 * inch:
            c.showPage()
            y = height - 1.0 * inch
            c.setFont("Helvetica", 10)

    y -= 0.2 * inch
    c.setFont("Helvetica-Bold", 12)
    c.drawString(1.0 * inch, y, "Budget Summary by Status")
    y -= 0.25 * inch
    c.setFont("Helvetica", 10)
    if summary.empty:
        c.drawString(1.0 * inch, y, "No summary available for current filters.")
    else:
        for _, r in summary.iterrows():
            s = (
                f"{r['Status']}: Students {int(r['Students']):,}, Budget "
                + ("-" if pd.isna(r["Budget"]) else f"${r['Budget']:,.0f}")
            )
            c.drawString(1.0 * inch, y, s)
            y -= 0.18 * inch
            if y < 1.0 * inch:
                c.showPage()
                y = height - 1.0 * inch
                c.setFont("Helvetica", 10)

    c.showPage()
    c.save()
    return bio.getvalue()

# =========================================================
# EXPORTS
# =========================================================
st.subheader("Exports")
colA, colB, colC = st.columns(3)

with colA:
    docx_bytes = build_docx_report()
    st.download_button(
        "Download Report (Word, .docx)",
        data=docx_bytes,
        file_name="EducationAbroad_Budget_Snapshot.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

with colB:
    pdf_bytes = build_pdf_report()
    st.download_button(
        "Download Report (PDF, .pdf)",
        data=pdf_bytes,
        file_name="EducationAbroad_Budget_Snapshot.pdf",
        mime="application/pdf",
    )

with colC:
    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="xlsxwriter") as writer:
        if not agg_confirmed.empty:
            agg_confirmed.to_excel(writer, index=False, sheet_name="Approved_Student_Budget")
        if not agg_pending.empty:
            agg_pending.to_excel(writer, index=False, sheet_name="Pending_Student_Budget")
        if not summary.empty:
            summary.to_excel(writer, index=False, sheet_name="Budget_Summary_by_Status")
    st.download_button(
        "Download Workbook (.xlsx)",
        data=out.getvalue(),
        file_name="EducationAbroad_Budget_Workbook.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
